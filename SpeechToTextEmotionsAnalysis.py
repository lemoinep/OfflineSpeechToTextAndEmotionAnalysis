# Author(s): Dr. Patrick Lemoine

# Offline speech-to-text + punctuation + Emotions Analysis
# - WAV Transcript -> TXT / JSON / DOCX
# - Silence-based line jumps
# - Restoration punctuation
# - Majuscules at the beginning of sentence / paragraph
# - Emotional analysis + colored DOCX

import os
import sys
import subprocess
import wave
import json
import argparse
import re
from collections import defaultdict

import numpy as np
from sklearn.cluster import KMeans

from vosk import Model, KaldiRecognizer, SpkModel, SetLogLevel
from tqdm import tqdm
from deepmultilingualpunctuation import PunctuationModel
from docx import Document
import matplotlib.pyplot as plt
import pandas as pd
from nrclex import NRCLex
import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt


SetLogLevel(-1)

EMOTION_SIGNS = {
    "fear": -1,
    "anger": -1,
    "anticip": 1,
    "trust": 1,
    "surprise": 1,
    "positive": 1,
    "negative": -1,
    "sadness": -1,
    "disgust": -1,
    "joy": 1,
}

EMOTION_COLORS_SINGLE = {
    "fear": WD_COLOR_INDEX.RED,
    "anger": WD_COLOR_INDEX.DARK_RED,
    "anticip": WD_COLOR_INDEX.GRAY_25,
    "trust": WD_COLOR_INDEX.TURQUOISE,
    "surprise": WD_COLOR_INDEX.PINK,
    "positive": WD_COLOR_INDEX.BRIGHT_GREEN,
    "negative": WD_COLOR_INDEX.RED,
    "sadness": WD_COLOR_INDEX.VIOLET,
    "disgust": WD_COLOR_INDEX.DARK_YELLOW,
    "joy": WD_COLOR_INDEX.YELLOW,
}


def get_coeff_emotion(top_emotions):
    content = str(top_emotions)
    coeff = 0
    if not top_emotions:
        return 0
    for emo, sign in EMOTION_SIGNS.items():
        if emo in content:
            coeff += sign
    return coeff


def ensure_output_dir(path):
    os.makedirs(path, exist_ok=True)
    return path


def analyze_text_file(path_txt, qview=False):
    with open(path_txt, encoding="utf-8") as f:
        text = f.read()

    if qview:
        print("\n\n", text, "\n\n")

    base_name = os.path.splitext(path_txt)[0]
    text_object = NRCLex(text)

    if qview:
        print("NbSentences =", len(text_object.sentences))

    affect_dict = text_object.affect_dict
    with open(base_name + "_Report_Analysis.csv", "w", encoding="utf-8") as f:
        for emo, value in affect_dict.items():
            if qview:
                print(emo, value)
            f.write(f"{emo},{value}\n")

    affect_freq = text_object.affect_frequencies
    with open(base_name + "_Report.csv", "w", encoding="utf-8") as f:
        f.write("Emotions,Frequencies\n")
        if qview:
            print("\nNb frequencies =", len(affect_freq))
        for emo, freq in affect_freq.items():
            if qview:
                print(emo, freq)
            f.write(f"{emo},{freq}\n")

    data = pd.read_csv(base_name + "_Report.csv")
    df = pd.DataFrame(data)

    X = list(df.iloc[:, 0])
    Y = list(df.iloc[:, 1])
    
    del X[2]
    del Y[2]
    
    #print(X)
    #print(Y)

    fig = plt.figure()
    max_y_lim = max(Y) + 0.01
    min_y_lim = min(Y)
    min_y_lim = 0.0
    plt.ylim(min_y_lim, max_y_lim)

    bars = plt.bar(X, Y)

    color_map = [
        "red",       # fear
        "red",       # anger
        #"gray",      # anticip
        "blue",      # trust
        "pink",      # surprise
        "green",     # positive
        "red",       # negative
        "red",       # sadness
        "red",       # disgust
        "yellow",    # joy
        "gray"       # anticipation
    ]

    for idx, bar in enumerate(bars):
        if idx < len(color_map):
            bar.set_color(color_map[idx])

    plt.title("Emotions Analysis")
    plt.xlabel("")
    plt.ylabel("Percentage")
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.savefig(base_name + "_Report.jpg")
    plt.close(fig)
    
    #
    plt.title("Emotions Analysis")
    #explode = (0.10, 0.11, 0.12, 0.13, 0.14, 0.15, 0.16, 0.17, 0.18, 0.19, 0.20)  
    explode = (0.10, 0.11, 0.12, 0.13, 0.14, 0.15, 0.16, 0.17, 0.18, 0.19)  
    fig, ax = plt.subplots(figsize=(10, 8))
    wedges, texts, autotexts = ax.pie(Y, explode=explode, labels=X, colors=color_map,
                                  autopct='%1.1f%%', pctdistance=0.85,
                                  shadow=True, startangle=90, textprops={'fontsize': 12})

    centre_circle = plt.Circle((0, 0), 0.75, fc='white', linewidth=0)
    fig.gca().add_artist(centre_circle)

    ax.axis('equal') 
    plt.title('Emotions Analysis', fontsize=16, fontweight='bold', pad=20)
    
    plt.tight_layout()
    plt.savefig(base_name + "_Pie_Report.jpg")
    plt.close(fig)
    #
    
    doc = docx.Document()
    doc.add_heading("Analysis of text emotions", 0)
    para = doc.add_paragraph("")

    font_styles = doc.styles
    if "CommentsStyle" in font_styles:
        font_charstyle = font_styles["CommentsStyle"]
    else:
        font_charstyle = font_styles.add_style("CommentsStyle", WD_STYLE_TYPE.CHARACTER)
        font_object = font_charstyle.font
        font_object.size = Pt(10)
        font_object.name = "Times New Roman"

    for sentence in text_object.sentences:
        phrase = f"{sentence} "
        phrase_object = NRCLex(phrase)
        top = phrase_object.top_emotions
        content = str(top)

        if len(top) == 1:
            color = WD_COLOR_INDEX.WHITE
            for emo, wd_color in EMOTION_COLORS_SINGLE.items():
                if emo in content:
                    color = wd_color
                    break
            run = para.add_run(phrase, style="CommentsStyle")
            run.font.highlight_color = color
        else:
            para.add_run(phrase, style="CommentsStyle").font.highlight_color = WD_COLOR_INDEX.WHITE

    doc.save(base_name + "_Report.docx")

    doc2 = docx.Document()
    doc2.add_heading("Analysis of text emotions", 0)
    para2 = doc2.add_paragraph("")

    font_styles2 = doc2.styles
    if "CommentsStyle" in font_styles2:
        font_charstyle2 = font_styles2["CommentsStyle"]
    else:
        font_charstyle2 = font_styles2.add_style("CommentsStyle", WD_STYLE_TYPE.CHARACTER)
        font_object2 = font_charstyle2.font
        font_object2.size = Pt(10)
        font_object2.name = "Times New Roman"

    for sentence in text_object.sentences:
        phrase = f"{sentence} "
        phrase_object = NRCLex(phrase)
        top = phrase_object.top_emotions
        num = get_coeff_emotion(top)

        if top:
            if num <= -4:
                color = WD_COLOR_INDEX.DARK_RED
            elif num <= -3:
                color = WD_COLOR_INDEX.RED
            elif num <= -1:
                color = WD_COLOR_INDEX.VIOLET
            elif num == 0:
                color = WD_COLOR_INDEX.WHITE
            elif num >= 4:
                color = WD_COLOR_INDEX.BRIGHT_GREEN
            elif num >= 3:
                color = WD_COLOR_INDEX.GREEN
            else:  # 1 or 2
                color = WD_COLOR_INDEX.GRAY_25
        else:
            color = WD_COLOR_INDEX.WHITE

        run = para2.add_run(phrase, style="CommentsStyle")
        run.font.highlight_color = color

    doc2.save(base_name + "_Report_Colors.docx")


def check_wav_format(wf: wave.Wave_read):
    if wf.getnchannels() != 1 or wf.getsampwidth() != 2 or wf.getcomptype() != "NONE":
        raise ValueError("Audio file must be mono 16-bit PCM WAV.")


def cosine_dist(x, y):
    nx = np.array(x)
    ny = np.array(y)
    return 1 - np.dot(nx, ny) / (np.linalg.norm(nx) * np.linalg.norm(ny))


def assign_speaker_id(spk_vector, speakers_db, thr=0.25):
    """
    speakers_db: list of {'id': int, 'vec': [...]}
    thr: distance threshold to decide new speaker
    """
    if spk_vector is None:
        return None

    if not speakers_db:
        speakers_db.append({"id": 1, "vec": spk_vector})
        return 1

    best_id = None
    best_dist = 10.0
    for sp in speakers_db:
        d = cosine_dist(sp["vec"], spk_vector)
        if d < best_dist:
            best_dist = d
            best_id = sp["id"]

    if best_dist > thr:
        new_id = max(sp["id"] for sp in speakers_db) + 1
        speakers_db.append({"id": new_id, "vec": spk_vector})
        return new_id
    else:
        return best_id


def recluster_speakers(speaker_vectors, target_speakers=2):
    """
    speaker_vectors: dict {raw_id: [np.array(...), ...]}
    return: dict {raw_id: new_id}  with new_id in [1..target_speakers]
    """
    if not speaker_vectors:
        return {}

    raw_ids = sorted(speaker_vectors.keys())
    emb_list = []
    for rid in raw_ids:
        v = np.stack(speaker_vectors[rid], axis=0)
        emb_list.append(v.mean(axis=0))
    emb = np.stack(emb_list, axis=0)

    k = min(target_speakers, emb.shape[0])
    if k <= 1:
        return {rid: 1 for rid in raw_ids}

    km = KMeans(n_clusters=k, random_state=0, n_init=10)
    labels = km.fit_predict(emb)

    raw_to_new = {}
    for rid, lab in zip(raw_ids, labels):
        raw_to_new[rid] = int(lab + 1)
    return raw_to_new


def collect_words_with_timestamps(vosk_results):
    words = []
    for res in vosk_results:
        if not isinstance(res, dict):
            continue
        spk_id = res.get("speaker_id")
        for w in res.get("result", []):
            if "word" in w and "start" in w and "end" in w:
                words.append(
                    {
                        "word": w["word"],
                        "start": float(w["start"]),
                        "end": float(w["end"]),
                        "speaker": spk_id,
                    }
                )
    return words


def build_text_with_linebreaks(words, silence_threshold=0.8):
    if not words:
        return ""

    lines = []
    current_line_words = []

    prev_end = words[0]["end"]
    current_line_words.append(words[0]["word"])

    for w in words[1:]:
        gap = w["start"] - prev_end
        if gap > silence_threshold:
            lines.append(" ".join(current_line_words))
            current_line_words = [w["word"]]
        else:
            current_line_words.append(w["word"])
        prev_end = w["end"]

    if current_line_words:
        lines.append(" ".join(current_line_words))

    return "\n".join(lines)


def build_text_with_linebreaks_and_speakers(words, silence_threshold=0.8):
    if not words:
        return ""

    lines = []
    current_line = []
    current_speaker = words[0].get("speaker")
    prev_end = words[0]["end"]

    for w in words:
        spk = w.get("speaker")
        gap = w["start"] - prev_end
        if gap > silence_threshold or spk != current_speaker:
            if current_line:
                prefix = f"Speaker {current_speaker}: " if current_speaker else ""
                lines.append(prefix + " ".join(current_line))
            current_line = [w["word"]]
            current_speaker = spk
        else:
            current_line.append(w["word"])
        prev_end = w["end"]

    if current_line:
        prefix = f"Speaker {current_speaker}: " if current_speaker else ""
        lines.append(prefix + " ".join(current_line))

    return "\n".join(lines)


def save_to_docx(text, docx_path, title=None):
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    for line in text.splitlines():
        if line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph("")
    doc.save(docx_path)


def capitalize_after_punctuation(text):
    lines = text.split("\n")
    new_lines = []

    for line in lines:
        s = line

        def cap_first(match):
            return match.group(1) + match.group(2).upper()

        s = re.sub(r"^(\s*)([a-z])", cap_first, s)

        def cap_after_punct(match):
            punct = match.group(1)
            spaces = match.group(2)
            letter = match.group(3)
            return punct + spaces + letter.upper()

        s = re.sub(r"([\.!?])(\s+)([a-z])", cap_after_punct, s)

        new_lines.append(s)

    return "\n".join(new_lines)


def transcribe(
    audio_file,
    model_path,
    out_txt,
    out_json,
    out_docx,
    out_txt_speakers,
    chunk_size,
    use_punctuation=True,
    silence_threshold=0.8,
    enable_speakers=False,
    target_speakers=2,
):

    if not os.path.isdir(model_path):
        raise FileNotFoundError(f"Model directory not found: {model_path}")

    if not os.path.isfile(audio_file):
        raise FileNotFoundError(f"Audio file not found: {audio_file}")

    model = Model(model_path)

    spk_model = None
    if enable_speakers:
        spk_model_path = os.path.join(os.path.dirname(model_path), "model-spk")
        if os.path.isdir(spk_model_path):
            spk_model = SpkModel(spk_model_path)
            print(f"Speaker model loaded from: {spk_model_path}")
        else:
            print("Warning: speaker model not found, disabling speaker labeling.")
            enable_speakers = False

    with wave.open(audio_file, "rb") as wf:
        check_wav_format(wf)

        recognizer = KaldiRecognizer(model, wf.getframerate())
        if spk_model is not None:
            recognizer.SetSpkModel(spk_model)
        recognizer.SetWords(True)

        total_frames = wf.getnframes()
        text_chunks = []
        raw_results = []
        speakers_db = []
        speaker_vectors = defaultdict(list)

        print(f"Processing '{audio_file}' ({total_frames} frames)...")

        with tqdm(total=total_frames, desc="Transcription", unit="frame") as pbar:
            while True:
                data = wf.readframes(chunk_size)
                if len(data) == 0:
                    break

                if recognizer.AcceptWaveform(data):
                    res = json.loads(recognizer.Result())

                    spk_id = None
                    if enable_speakers and "spk" in res:
                        spk_vec = np.array(res["spk"], dtype=float)
                        spk_id = assign_speaker_id(spk_vec, speakers_db)
                        res["speaker_id"] = spk_id
                        speaker_vectors[spk_id].append(spk_vec)

                    raw_results.append(res)
                    text_chunks.append(res.get("text", ""))

                pbar.update(chunk_size)

        final_res = json.loads(recognizer.FinalResult())
        if enable_speakers and "spk" in final_res:
            spk_vec = np.array(final_res["spk"], dtype=float)
            spk_id = assign_speaker_id(spk_vec, speakers_db)
            final_res["speaker_id"] = spk_id
            speaker_vectors[spk_id].append(spk_vec)

        raw_results.append(final_res)
        text_chunks.append(final_res.get("text", ""))

    full_text = " ".join(chunk for chunk in text_chunks if chunk).strip()

    words = collect_words_with_timestamps(raw_results)

    new_speaker_map = {}
    if enable_speakers and speaker_vectors:
        new_speaker_map = recluster_speakers(speaker_vectors, target_speakers=target_speakers)

    if new_speaker_map:
        for res in raw_results:
            if isinstance(res, dict) and "speaker_id" in res:
                old_id = res["speaker_id"]
                if old_id in new_speaker_map:
                    res["speaker_id"] = new_speaker_map[old_id]

        for w in words:
            old_spk = w.get("speaker")
            if old_spk in new_speaker_map:
                w["speaker"] = new_speaker_map[old_spk]

    if enable_speakers:
        text_with_linebreaks_speakers = build_text_with_linebreaks_and_speakers(
            words, silence_threshold=silence_threshold
        )
        text_with_linebreaks = build_text_with_linebreaks(
            words, silence_threshold=silence_threshold
        )
    else:
        text_with_linebreaks = build_text_with_linebreaks(
            words, silence_threshold=silence_threshold
        )
        text_with_linebreaks_speakers = text_with_linebreaks

    punctuated_text = text_with_linebreaks
    punctuated_text_speakers = text_with_linebreaks_speakers

    if use_punctuation and text_with_linebreaks:
        print("\nRestoring punctuation...")
        punct_model = PunctuationModel()

        lines = text_with_linebreaks.split("\n")
        new_lines = []
        for line in lines:
            if line.strip():
                new_lines.append(punct_model.restore_punctuation(line))
            else:
                new_lines.append("")
        punctuated_text = "\n".join(new_lines)

        lines_spk = text_with_linebreaks_speakers.split("\n")
        new_lines_spk = []
        for line in lines_spk:
            if line.strip():
                new_lines_spk.append(punct_model.restore_punctuation(line))
            else:
                new_lines_spk.append("")
        punctuated_text_speakers = "\n".join(new_lines_spk)

    punctuated_text = capitalize_after_punctuation(punctuated_text)
    punctuated_text_speakers = capitalize_after_punctuation(punctuated_text_speakers)

    if (False):
        print("\n=== Raw transcription (no line breaks) ===")
        print(full_text)
        print("\n=== With line breaks (silence-based) ===")
        print(text_with_linebreaks)
        if enable_speakers:
            print("\n=== With line breaks + speakers ===")
            print(text_with_linebreaks_speakers)
    
        print("\n=== With line breaks + punctuation + caps (no speakers) ===")
        print(punctuated_text)
    
        if enable_speakers:
            print("\n=== With line breaks + speakers + punctuation + caps ===")
            print(punctuated_text_speakers)
        print("=====================")

    with open(out_txt, "w", encoding="utf-8") as f_txt:
        f_txt.write(punctuated_text + "\n")

    with open(out_txt_speakers, "w", encoding="utf-8") as f_txt_spk:
        f_txt_spk.write(punctuated_text_speakers + "\n")

    output_json_struct = {
        "results": raw_results,
        "words": words,
    }
    with open(out_json, "w", encoding="utf-8") as f_json:
        json.dump(output_json_struct, f_json, ensure_ascii=False, indent=2)

    save_to_docx(
        punctuated_text,
        out_docx,
        title=f"Transcription - {os.path.basename(audio_file)}",
    )



if __name__ == "__main__":

    parser = argparse.ArgumentParser(
        description=(
            "Offline speech-to-text with Vosk "
            "(progress bar, silence-based line breaks, punctuation, TXT + DOCX output, "
            "optional speaker labeling + reclustering, emotion analysis)."
        )
    )

    parser.add_argument(
        "--Path", type=str, default=".", help="Path to WAV file directory"
    )

    parser.add_argument(
        "--Name",
        type=str,
        default="audio.wav",
        help="Name of mono 16 kHz WAV file (or mp3/mp4 to convert).",
    )

    parser.add_argument(
        "-m",
        "--Model",
        type=str,
        default="vosk-model-small-en-us-0.15",
        help="Path to Vosk model directory (relative to ./Models).",
    )

    parser.add_argument(
        "-c",
        "--chunk_size",
        type=int,
        default=4000,
        help="Chunk size (frames) for reading the WAV file (default: 4000).",
    )

    parser.add_argument(
        "--no-punct",
        action="store_true",
        help="Disable punctuation restoration (keep only silence-based line breaks).",
    )

    parser.add_argument(
        "--emotions_analysis",
        type=int,
        default=1,
        help="Enable Emotion Analysis mode",
    )

    parser.add_argument(
        "--silence_threshold",
        type=float,
        default=0.8,
        help=(
            "Silence threshold in seconds to insert a line break between words "
            "(default: 0.8)."
        ),
    )

    parser.add_argument(
        "--enable_speakers",
        action="store_true",
        help="Enable simple speaker labeling using Vosk speaker model.",
    )

    parser.add_argument(
        "--target_speakers",
        type=int,
        default=2,
        help="Number of main speakers to recluster to (default: 2).",
    )

    args = parser.parse_args()
    
    qtranscribe = True

    PathW = os.path.dirname(sys.argv[0])
    MODEL_PATH = os.path.join(PathW, "Models", args.Model)

    input_file = os.path.join(args.Path, args.Name)
    base_name, ext = os.path.splitext(os.path.basename(args.Name))  # ext = .wav / .mp3 / .mp4...[web:4]

    final_wav = input_file

    if ext.lower() in [".mp4", ".mp3"]:
        final_wav = os.path.join(args.Path, f"{base_name}.wav")

        # Commande ffmpeg
        cmd = [
            "ffmpeg",
            "-y",              # overwrite
            "-i", input_file,  # input
            "-vn",             
            "-ac", "1",        # mono
            "-ar", "16000",    # 16 kHz
            "-acodec", "pcm_s16le",  # PCM 16 bits little-endian
            final_wav,
        ]

        result = subprocess.run(cmd)
        if result.returncode != 0:
            print("Erreur lors de la conversion ffmpeg.")
            sys.exit(1)

    if ext.lower() in [".txt"]:
        qtranscribe = False
        
    AUDIO_FILE = final_wav

    output_dir = os.path.join(args.Path, "DATA")
    os.makedirs(output_dir, exist_ok=True)

    base_name = os.path.splitext(os.path.basename(AUDIO_FILE))[0]
    OutputJson = os.path.join(output_dir, f"{base_name}.json")
    OutputTxt = os.path.join(output_dir, f"{base_name}.txt")
    OutputTxtSpeakers = os.path.join(output_dir, f"{base_name}_speakers.txt")
    OutputDocx = os.path.join(output_dir, f"{base_name}.docx")

    use_punct = not args.no_punct
    
    if qtranscribe :
        transcribe(
            audio_file=AUDIO_FILE,
            model_path=MODEL_PATH,
            out_txt=OutputTxt,
            out_json=OutputJson,
            out_docx=OutputDocx,
            out_txt_speakers=OutputTxtSpeakers,
            chunk_size=args.chunk_size,
            use_punctuation=use_punct,
            silence_threshold=args.silence_threshold,
            enable_speakers=args.enable_speakers,
            target_speakers=args.target_speakers,
        )

    if args.emotions_analysis:
        analyze_text_file(OutputTxt, qview=False)

    print("\n--- Finished ---")
